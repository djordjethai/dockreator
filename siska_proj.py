import os
import re
import json
import fitz
import pandas as pd
from docx import Document
from openai import OpenAI
import streamlit as st
import zipfile
import datetime

client = OpenAI()

def extract_keys_from_template(template_docx_path):
    doc = Document(template_docx_path)
    template_text = ""

    # Collect text from paragraphs
    for paragraph in doc.paragraphs:
        template_text += paragraph.text + "\n"

    # Collect text from tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                template_text += cell.text + "\n"

    # Regex to find keys enclosed in double curly braces
    keys = re.findall(r'\[([^\]]+)\]', template_text)
    return keys

def read_txt_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as file:
        return file.read()

def read_docx_file(filepath):
    doc = Document(filepath)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    return '\n'.join(text)

def read_pdf_file(filepath):
    doc = fitz.open(filepath)
    text = ""
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    return text

def read_documents_from_folders(folders):
    documents = {}
    for folder_path in folders:
        for filename in os.listdir(folder_path):
            filepath = os.path.join(folder_path, filename)
            if filename.endswith('.txt'):
                documents[filename] = read_txt_file(filepath)
            elif filename.endswith('.docx'):
                documents[filename] = read_docx_file(filepath)
            elif filename.endswith('.pdf'):
                documents[filename] = read_pdf_file(filepath)
    return documents

def extract_values_from_csv(csv_path, member, opcija):
    df = pd.read_csv(csv_path)
    row = df[(df['member'] == member) & (df['opcija'] == opcija)]
    if not row.empty:
        return row.iloc[0].to_dict()
    return {}

def extract_values_from_docs(keys, documents, csv_values):
    response_text = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "user",
                "content": f"""The following text contains values for the following keys: {keys}. \
                Please identify the values for those entities in {documents} and extract the values, \
                and return key value pairs formatted as JSON.\n."""
            }
        ],
        temperature=0,
        response_format={"type": "json_object"},
    )
    llm_response = response_text.choices[0].message.content
    values_dict = json.loads(llm_response)
    
    # Filter the values_dict to only include the keys we are interested in
    filtered_values_dict = {key: values_dict[key] for key in keys if key in values_dict}
    
    # Add the CSV values to the filtered values dictionary
    filtered_values_dict.update(csv_values)
    
    with st.expander("Pronadjeni podaci", expanded=True):
        st.write(filtered_values_dict)
    return filtered_values_dict

def create_filled_template(template_docx_path, values_dict, output_docx_path):
    doc = Document(template_docx_path)
    for paragraph in doc.paragraphs:
        for key, value in values_dict.items():
            if f'[{key}]' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'[{key}]', value)
    doc.save(output_docx_path)

def zip_specific_files(files_to_zip, zip_name):
  
    with zipfile.ZipFile(zip_name, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file_path in files_to_zip:
            if os.path.exists(file_path):
                zipf.write(file_path, os.path.basename(file_path))
       

def main(template_path, folders, base_folder, fixed_folder, members, opcija, podela):
    keys = extract_keys_from_template(template_path)
    files_to_zip = []
    for member in members:
        st.info(f"Kreiram {opcija} za {member}")
        # Create a new list of folders with the member's name added to each base folder
        updated_folders = [os.path.join(folder, member) for folder in folders]
        updated_folders.append(fixed_folder)
        # Read documents from the updated list of folders
        documents = read_documents_from_folders(updated_folders)
        st.info(f"Čitam dokumente: {list(documents.keys())}")
        output_docx_path = os.path.join(base_folder, member, f"{member} {opcija}.docx")
        csv_values = extract_values_from_csv(podela, member, opcija)
        values_dict = extract_values_from_docs(keys, documents, csv_values)
        create_filled_template(template_path, values_dict, output_docx_path)
        files_to_zip.append(output_docx_path)
    zip_specific_files(files_to_zip, f"{opcija}.zip")
    st.success("Kreirani su svi zahtevani dokumenti")

if __name__ == "__main__":
    template_path = './zasisku/mustre/tpl.docx'  # Ensure your template uses delimiters [key]
    base_folders = ['./zasisku/zaposleni']
    base_folder = './zasisku/zaposleni'
    fixed_folder = './zasisku/dokumenti'
    available_members = [name for name in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, name))]
    folder_path = './zasisku' 
    podela = './zasisku/podela/podela.csv'
    st.title("Siska - Kreiranje dokumenata za zaposlene")    
    st.subheader("Kreiranje dokumenata za zaposlene")
    st.caption("Ver 11.06.2024.")
    if "gotovo" not in st.session_state:
        st.session_state.gotovo = False
        
    with st.form("my_form"):
        opcija = st.selectbox("Izaberite tip dokumenta", ["Ugovor o radu", "Preciscen tekst ugovora o radu", "Aneks Ugovora o radu", "Izjava o poverljivosti informacija", "Rešenje o godišnjem odmoru"], index=None)
        members = st.multiselect("Odaberite zaposlene", available_members)
        if st.form_submit_button("Kreiraj dokumente"):
            main(template_path, base_folders, base_folder, fixed_folder, members, opcija, podela)
            st.info("Kreirani su svi zahtevani dokumenti")
            st.session_state.gotovo = True
    if st.session_state.gotovo:
        with open(f"{opcija}.zip", "rb") as fp:
            btn = st.download_button(
                label="Preuzmi ZIP fajl",
                data=fp,
                file_name=f"{opcija}.zip",
                mime="application/zip"
            )
