from docx import Document
from openai import OpenAI
import os

# Initialize OpenAI client
client = OpenAI()
conversation_history = []
finish_reason = ""


# X:\05 OPŠTI\LJUDSKI RESURSI\RADNICI\Bojan Gavrilović\UGOVOR O RADU\KONACNO
# "C:\\Users\\djordje\\Desktop\\RADNICI\\Bojan Gavrilović\\UGOVOR O RADU\\KONACNO"

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    print(f"procitao tekst {file_path}")    
    return '\n'.join(full_text)

def add_markdown_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    bold = False
    parts = text.split('**')
    for part in parts:
        run = p.add_run(part)
        if bold:
            run.bold = True
        bold = not bold


def save_to_docx(template_file, save_path, final_text):
    doc = Document(template_file)
    # Clear existing content
    doc._body.clear_content()
    lines = final_text.split('\n')
    print(f"Cuvam tekst u {save_path}")
    # Add the final text, handling markdown-like formatting
    for line in lines:
        if line.startswith('# '):
            add_markdown_paragraph(doc, line[2:], style='Heading 1')
        elif line.startswith('## '):
            add_markdown_paragraph(doc, line[3:], style='Heading 2')
        elif line.startswith('### '):
            add_markdown_paragraph(doc, line[4:], style='Heading 3')
        elif line.startswith('#### '):
            add_markdown_paragraph(doc, line[5:], style='Heading 4')
        elif line.startswith('- ') or line.startswith(' - ') or line.startswith('  - ') or line.startswith('   - ') or line.startswith(' -  ') or line.startswith('  -  '):
            add_markdown_paragraph(doc, line[2:], style='List Paragraph')
        else:    
            add_markdown_paragraph(doc, line)
    doc.save(save_path)

def refine_with_llm(contract_text, annex_text, finish_reason):
    # Using OpenAI's GPT to ensure coherence and integration
    story = ""
    while True:
        if finish_reason == "" or finish_reason == None:
            conversation_history.append({"role": "user", "content": f"""Please update the following contract by replacing the existing sentences with the corresponding new sentences from the annex. Ensure that all changes are clearly integrated and the document remains coherent and readable. Use appropriate headings (H2, H3, etc.) to organize the text for better readability.

                    Original Contract:
                    {contract_text}

                    Annex with Changes:
                    {annex_text}

                    Instructions:
                    1. Identify and replace sentences in the original contract with the corresponding new sentences from the annex.
                    2. Maintain the structure and coherence of the contract.
                    3. Do not give the title for the document.
                    4. Use appropriate formatting such as H2 for articles, and so on to organize the text.
                    5. Leave out sentence like 'menja se i glasi' and the quotation marks associated with it.

                    Output the updated contract below:
                    """})
        elif finish_reason == "length":
            conversation_history.append({"role": "user", "content": "Continue"})
        print("Obradjujem tekst...")
        response = client.chat.completions.create(
            model="gpt-4o",
            temperature=0,
            messages=conversation_history,
            stop=None,
        )
        conversation_history.append({"role": "assistant", "content": response.choices[0].message.content})
        finish_reason = response.choices[0].finish_reason
        print(finish_reason)
        story += response.choices[0].message.content
        
        if finish_reason == "stop":
            return story

def process_subfolders(base_folder):
    i=0 # ovo je samo za test
    # Loop through each person’s folder in the base folder
    for person_name in os.listdir(base_folder):
        i+=1
        person_folder = os.path.join(base_folder, person_name)

        if os.path.isdir(person_folder) and not "!" in person_name and i<5:
            
            print("Radim", person_name)
            ugovor_o_radu_folder = os.path.join(person_folder, 'UGOVOR O RADU', 'KONACNO')
            if os.path.exists(ugovor_o_radu_folder):
                contract_file = None
                annex_file = None

                # Assign the correct files to contract_file and annex_file
                for file in os.listdir(ugovor_o_radu_folder):
                    if file.startswith('Ugovor'):
                        contract_file = os.path.join(ugovor_o_radu_folder, file)
                    elif file.startswith('Aneks 1'):
                        annex_file = os.path.join(ugovor_o_radu_folder, file)
                
                if contract_file and annex_file:
                    template_file = 'template.docx'  # Fixed template file
                    output_file = os.path.join(ugovor_o_radu_folder, f'Preciscen_{os.path.basename(contract_file)}')
                    
                    # Step 1: Read DOCX files
                    contract_text = read_docx(contract_file)
                    annex_text = read_docx(annex_file)
                    
                    # Step 3: Refine the aggregated text using LLM (optional)
                    finish_reason = None  # Define your finish_reason as needed
                    final_text = refine_with_llm(contract_text, annex_text, finish_reason)

                    # Step 4: Save the refined text to a new DOCX file created from a template
                    save_to_docx(template_file, output_file, final_text)
                    print(f"Processed: {ugovor_o_radu_folder}")

# Define the base folder path
base_folder = 'C:\\Users\\djordje\\Desktop\\RADNICI'
process_subfolders(base_folder)
